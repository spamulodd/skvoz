#!/usr/bin/env python3
"""Generate a one-page WiFi QR DOCX. Credentials via CLI only (never commit secrets)."""
from __future__ import annotations

import argparse
import re
from io import BytesIO
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DEFAULT = Path(__file__).with_name("wifi-qr.docx")


def esc_wifi(value: str) -> str:
    return re.sub(r'([\\;,:"])', r"\\\1", value)


def set_run_font(run, name="Calibri", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_paragraph(paragraph, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    paragraph._p.get_or_add_pPr().append(shd)


def add_centered(doc, text, *, size=12, bold=False, color=None, space_before=0, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def build_docx(ssid: str, password: str, out: Path) -> None:
    payload = f"WIFI:T:WPA;S:{esc_wifi(ssid)};P:{esc_wifi(password)};;"
    qr = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=12,
        border=2,
    )
    qr.add_data(payload)
    qr.make(fit=True)
    img = qr.make_image(fill_color="#0f1a14", back_color="white")
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    add_centered(doc, "RVPN · OpenWrt", size=11, bold=True, color=(31, 107, 74), space_after=6)
    add_centered(doc, "Новый WiFi с VPN", size=28, bold=True, color=(15, 26, 20), space_after=10)
    add_centered(
        doc,
        "Отсканируйте QR-код камерой телефона — сеть подключится автоматически.",
        size=12,
        color=(61, 82, 72),
        space_after=10,
    )

    notice = add_centered(
        doc,
        "Работает в тестовом режиме. Возможны сайты, которые не открываются или открываются нестабильно.",
        size=11,
        bold=True,
        color=(140, 80, 20),
        space_before=4,
        space_after=18,
    )
    shade_paragraph(notice, "FFF4E5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(buf, width=Cm(8.5))

    add_centered(doc, "Сканировать для подключения", size=10, color=(61, 82, 72), space_after=22)
    add_centered(doc, "СЕТЬ", size=10, bold=True, color=(61, 82, 72), space_after=4)
    add_centered(doc, ssid, size=18, bold=True, color=(15, 26, 20), space_after=14)
    add_centered(doc, "ПАРОЛЬ", size=10, bold=True, color=(61, 82, 72), space_after=4)
    add_centered(doc, password, size=18, bold=True, color=(15, 26, 20), space_after=14)
    add_centered(doc, "2.4 GHz · 5 GHz · WPA/WPA2", size=10, color=(61, 82, 72), space_after=0)

    doc.save(out)
    print(f"saved {out}")


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate WiFi QR one-pager (DOCX)")
    ap.add_argument("--ssid", required=True, help="Wi-Fi network name")
    ap.add_argument("--password", required=True, help="Wi-Fi password")
    ap.add_argument("-o", "--output", type=Path, default=OUT_DEFAULT, help="Output .docx path")
    args = ap.parse_args()
    build_docx(args.ssid, args.password, args.output)


if __name__ == "__main__":
    main()
